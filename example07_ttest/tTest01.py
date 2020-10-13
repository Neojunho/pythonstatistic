import pandas as pd
import numpy as np
from scipy import stats
import matplotlib.pyplot as plt
import docx
import os

# RawData 로드하기
data = pd.read_csv("./example07_ttest/ttestData.csv", na_values=[999]) # RawData 파일 경로와 결측치 값 설정
df = pd.DataFrame(data, columns = ['성별', '스마트폰중독1', '스마트폰중독2', '스마트폰중독3']) # 컬럼 정의
df = df.dropna() # 결측치 column 제거

# 스마트폰중독 3문항에 대한 평균 구하기
df["스마트폰중독평균"] = (df["스마트폰중독1"] + df["스마트폰중독2"] + df["스마트폰중독3"])/3

# numpy로 남여 두 그룹을 나누는 배열 선언
group1 = np.array([])
group2 = np.array([])
arrayIndex = 0

# 성별로 집단 구분하기
for i in df["성별"]:
    if(i == 1):
        group1 = np.append(group1, np.array(df["스마트폰중독평균"][arrayIndex]))
    else :
        group2 = np.append(group2, np.array(df["스마트폰중독평균"][arrayIndex]))
    arrayIndex += 1

#Levene의 등분산 검정 
lResult = stats.levene(group1, group2,  center="mean")
print('LeveneResult(F) : %.3f \np-value : %.3f' % (lResult))
leveneF = lResult[0]
leveneP = lResult[1]

남자사례수 = len(group1)
print("남자사례수", 남자사례수)
여자사례수 =len(group2)
print("여자사례수", 여자사례수)
남자평균 = np.mean(group1)
print("남자평균:", 남자평균)
남자분산 = np.var(group1)
print("남자분산:", 남자분산)
남자표준편차 = np.std(group1)
print("남자표준편차:", 남자표준편차)
여자평균 = np.mean(group2)
print("여자평균:", 여자평균)
여자분산 = np.var(group2)
print("여자분산:", 여자분산)
여자표준편차 = np.std(group2)
print("여자표준편차:", 여자표준편차)

tResult_equal = stats.ttest_ind(group1,group2)
print('등분산 가정됨 t-test(T): %.3f \np-value : %.3f' % (tResult_equal))
등분산t = tResult_equal[0]
등분사p = tResult_equal[1]

tResult_NonEqual = stats.ttest_ind(group1,group2, equal_var = False)
print('등분산 가정안됨 t-test(T): %.3f \np-value : %.3f' % (tResult_NonEqual))
등분산아님t = tResult_NonEqual[0]
등분산아님p = tResult_NonEqual[1]

doc = docx.Document()
doc.add_heading('<표 1> 스마트폰중독 남여 차이 비교', 0)

records = [
	['스마트폰중독', '남자', 남자사례수, 남자평균, 남자표준편차, 등분산t, 등분사p],
	['스마트폰중독', '여자', 여자사례수, 여자평균, 여자표준편차, 등분산t, 등분사p]
		   ]

menuTable = doc.add_table(rows=1,cols=7)
menuTable.style= 'Medium Shading 2 Accent 3'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = '종속변수'
hdr_Cells[1].text = '성별'
hdr_Cells[2].text = '사례수(n)'
hdr_Cells[3].text = '평균(M)'
hdr_Cells[4].text = '표준편차(SD)'
hdr_Cells[5].text = 't'
hdr_Cells[6].text = 'p'

for col1, col2, col3, col4, col5, col6, col7 in records:
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text = col1
    row_Cells[1].text = col2
    row_Cells[2].text = str(round(col3,3))
    row_Cells[3].text = str(round(col4,3))
    row_Cells[4].text = str(round(col5,3))
    row_Cells[5].text = str(round(col6,3))
    row_Cells[6].text = str(round(col7,3))

doc.save('tTestResult.docx')
os.system("start tTestResult.docx")